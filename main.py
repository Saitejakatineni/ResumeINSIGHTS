"""
resumeINSIGHTS — FastAPI Backend
Run: uvicorn main:app --reload --port 8000
Or launch as desktop app: python app.py
"""

import base64
import json
import os
import re
import sqlite3
import tempfile
from datetime import datetime

import anthropic
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles

# Support both dev (.env in project root) and bundled (~/Library/Application Support)
_env_file = os.environ.get('RESUMEINSIGHTS_ENV', os.path.join(os.path.dirname(__file__), '.env'))
load_dotenv(_env_file)

app = FastAPI(title="resumeINSIGHTS")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

client = anthropic.Anthropic()

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

DB_PATH = os.environ.get(
    'RESUMEINSIGHTS_DB',
    os.path.join(os.path.dirname(__file__), "history.db")
)
STATIC_DIR = os.environ.get(
    'RESUMEINSIGHTS_STATIC',
    os.path.join(os.path.dirname(__file__), "static")
)


# ── Database ──────────────────────────────────────────────────────────────────

def init_db():
    con = sqlite3.connect(DB_PATH)
    con.execute("""
        CREATE TABLE IF NOT EXISTS analyses (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at  TEXT NOT NULL,
            resume_name TEXT NOT NULL,
            jd_snippet  TEXT NOT NULL,
            jd_full     TEXT NOT NULL,
            result_json TEXT NOT NULL
        )
    """)
    con.commit()
    con.close()


def save_analysis(resume_name: str, jd_full: str, result: dict) -> int:
    con = sqlite3.connect(DB_PATH)
    cur = con.execute(
        "INSERT INTO analyses (created_at, resume_name, jd_snippet, jd_full, result_json) VALUES (?,?,?,?,?)",
        (
            datetime.now().strftime("%Y-%m-%d %H:%M"),
            resume_name,
            jd_full[:120].strip(),
            jd_full,
            json.dumps(result),
        ),
    )
    row_id = cur.lastrowid
    con.commit()
    con.close()
    return row_id


def get_history():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    rows = con.execute(
        "SELECT id, created_at, resume_name, jd_snippet, result_json FROM analyses ORDER BY id DESC LIMIT 50"
    ).fetchall()
    con.close()
    return [
        {
            "id": r["id"],
            "created_at": r["created_at"],
            "resume_name": r["resume_name"],
            "jd_snippet": r["jd_snippet"],
            "match_percentage": json.loads(r["result_json"]).get("match_percentage"),
            "ats_score": json.loads(r["result_json"]).get("ats_score"),
            "overall_grade": json.loads(r["result_json"]).get("overall_grade"),
        }
        for r in rows
    ]


def get_analysis_by_id(analysis_id: int):
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    row = con.execute("SELECT * FROM analyses WHERE id = ?", (analysis_id,)).fetchone()
    con.close()
    if not row:
        return None
    return {
        "id": row["id"],
        "created_at": row["created_at"],
        "resume_name": row["resume_name"],
        "jd_full": row["jd_full"],
        "result": json.loads(row["result_json"]),
    }


init_db()


# ── Helpers ───────────────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        doc = docx.Document(tmp_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    finally:
        os.unlink(tmp_path)


ANALYSIS_PROMPT = """Your task: analyze this resume against the job description and give the candidate everything they need to reach a 97%+ ATS match. Be brutally honest and hyper-specific — every suggestion must cite something concrete from the resume or JD. Generic advice is useless here.

<job_description>
{job_description}
</job_description>

Think carefully before answering. Work through these dimensions:

1. KEYWORD ANALYSIS — scan the JD for every skill, tool, methodology, certification, and domain term. Mark each as present or absent in the resume. Exact phrasing matters for ATS parsers.

2. SKILLS & EXPERIENCE FIT — do the candidate's years of experience, seniority level, technical stack, and domain match what the JD requires? Note any mismatches honestly.

3. IMPACT & LANGUAGE — are accomplishments quantified with metrics? Do bullet points start with strong action verbs? Does the language mirror the JD's tone (e.g. "led" vs "managed")?

4. ATS PARSEABILITY — would an ATS correctly extract job titles, dates, skills, and education? Flag tables, columns, headers, graphics, or unusual formatting that breaks parsers.

5. GAP-CLOSING SUGGESTIONS — for every missing keyword or weak area, provide a concrete, ready-to-paste rewrite or new bullet point. The improved lines must sound natural, not keyword-stuffed.

Return ONLY a valid JSON object — no explanation, no markdown fences, just raw JSON:
{{
  "match_percentage": <integer 0-100, based purely on skills+experience overlap with JD requirements>,
  "ats_score": <integer 0-100, based on keyword density, formatting, and parseability>,
  "overall_grade": "<A+|A|A-|B+|B|B-|C+|C|C-|D|F>",
  "match_summary": "<3-4 sentences: current fit level, biggest strengths, biggest gaps, realistic chance of passing ATS>",
  "matching_keywords": ["<exact keyword or phrase found in both resume and JD>"],
  "missing_keywords": ["<high-value JD keyword completely absent from resume — sort by importance>"],
  "strengths": ["<specific strength with direct evidence from the resume, tied to this JD>"],
  "weaknesses": ["<specific gap with direct reference to what the JD requires that is missing>"],
  "lines_to_change": [
    {{
      "original": "<copy the existing resume line verbatim or closely paraphrased>",
      "improved": "<rewritten version: stronger verb, quantified, JD keywords woven in naturally>",
      "reason": "<exactly which JD requirement this targets and why the new phrasing scores higher>"
    }}
  ],
  "lines_to_add": [
    {{
      "line": "<complete, ready-to-paste bullet point or sentence — not a template, a real line>",
      "section": "<Summary|Experience|Skills|Education|Projects|Certifications>",
      "reason": "<which missing JD keyword or requirement this addresses>"
    }}
  ],
  "formatting_tips": ["<specific, actionable tip — name the exact element to fix and how>"],
  "priority_actions": [
    "<#1 single highest-impact change that will most improve the ATS score>",
    "<#2 second highest-impact change>",
    "<#3 third highest-impact change>"
  ]
}}

Scoring calibration:
- match_percentage 90-100: candidate meets almost all requirements, would likely pass recruiter screen
- match_percentage 70-89: solid fit with addressable gaps
- match_percentage 50-69: partial fit, significant rework needed
- match_percentage <50: major mismatch in experience or skills
- ats_score 90-100: clean formatting, high keyword density, all sections machine-readable
- ats_score 70-89: minor formatting issues or some keyword gaps
- ats_score <70: formatting or keyword problems likely to cause ATS rejection
"""


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/test-analysis")
async def test_analysis():
    """Returns a realistic fake result — no API call, no credits used."""
    result = {
        "match_percentage": 72,
        "ats_score": 68,
        "overall_grade": "B-",
        "match_summary": "The candidate has a solid Python and data engineering background that partially aligns with the role. Key cloud and pipeline tools from the JD (DBT, Snowflake, Airflow) are absent from the resume, which will likely cause ATS filtering. With targeted additions the match can reach 90%+.",
        "matching_keywords": ["Python", "SQL", "AWS", "ETL", "REST APIs", "CI/CD", "Agile", "data pipelines", "PostgreSQL"],
        "missing_keywords": ["DBT", "Snowflake", "Airflow", "MongoDB", "S3", "cross-functional collaboration", "iterative delivery"],
        "strengths": [
            "Strong Python experience (3+ years) directly matches the JD's primary language requirement.",
            "AWS experience aligns with the JD's cloud infrastructure expectations.",
            "Demonstrated ETL pipeline work covers a core responsibility of the role."
        ],
        "weaknesses": [
            "No mention of DBT or Snowflake — both are explicitly listed as 'a plus' in the JD and are common ATS filters.",
            "Missing quantified impact: JD asks for 'proven ability to deliver high-quality, iterative value' but resume bullets lack metrics.",
            "MongoDB not listed in skills section despite JD calling it out specifically."
        ],
        "lines_to_change": [
            {
                "original": "Built data pipelines to process customer data using Python and SQL.",
                "improved": "Designed and maintained Python-based ETL pipelines processing 10M+ daily records, reducing data latency by 40% using AWS S3 and PostgreSQL.",
                "reason": "Adds quantified impact and AWS S3 keyword the JD requires; 'designed and maintained' is a stronger verb pair than 'built'."
            },
            {
                "original": "Worked with the data team on reporting dashboards.",
                "improved": "Collaborated cross-functionally with data, product, and engineering teams to deliver iterative reporting solutions, cutting dashboard load time by 30%.",
                "reason": "Directly mirrors JD language: 'cross-functional collaboration' and 'iterative value' — both flagged by ATS."
            }
        ],
        "lines_to_add": [
            {
                "line": "Proficient in DBT (data build tool) for SQL-based data transformations and Snowflake for cloud data warehousing.",
                "section": "Skills",
                "reason": "DBT and Snowflake are explicitly listed in the JD and completely absent — highest priority ATS keyword gap."
            },
            {
                "line": "Experience with MongoDB for document storage and Apache Airflow for workflow orchestration in data pipeline environments.",
                "section": "Skills",
                "reason": "Closes two JD keyword gaps (MongoDB, modern pipeline tooling) in a single natural-sounding line."
            }
        ],
        "formatting_tips": [
            "Move Skills section above Experience — ATS parsers score keyword density in the first 30% of the document higher.",
            "Replace any multi-column layout with single-column — column-based PDFs often parse as garbled text in ATS systems.",
            "Ensure job titles exactly match industry-standard titles (e.g. 'Data Engineer' not 'Data Pipeline Developer') so ATS title-matching works."
        ],
        "priority_actions": [
            "Add DBT and Snowflake to the Skills section immediately — these are the #1 ATS filter keywords missing from your resume.",
            "Quantify at least 3 existing bullet points with metrics (rows processed, latency reduced, time saved) to satisfy the JD's 'proven delivery' requirement.",
            "Rewrite 2 experience bullets to include 'cross-functional' and 'iterative' phrasing to mirror the JD's collaboration language."
        ],
        "_usage": {"input_tokens": 0, "output_tokens": 0, "est_cost_usd": 0.0},
        "_test_mode": True
    }
    row_id = save_analysis("sample_resume.pdf", "Sample job description for testing.", result)
    result["history_id"] = row_id
    return JSONResponse(result)


@app.get("/history")
async def list_history():
    return JSONResponse(get_history())


@app.get("/history/{analysis_id}")
async def get_history_item(analysis_id: int):
    item = get_analysis_by_id(analysis_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found.")
    return JSONResponse(item)


@app.delete("/history/{analysis_id}")
async def delete_history_item(analysis_id: int):
    con = sqlite3.connect(DB_PATH)
    con.execute("DELETE FROM analyses WHERE id = ?", (analysis_id,))
    con.commit()
    con.close()
    return JSONResponse({"ok": True})


@app.post("/analyze")
async def analyze_resume(
    resume: UploadFile = File(...),
    job_description: str = Form(...),
):
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required.")

    content = await resume.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Max size is 10 MB.")

    filename = (resume.filename or "upload").lower()

    try:
        if filename.endswith(".pdf"):
            pdf_b64 = base64.standard_b64encode(content).decode("utf-8")
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_b64,
                            },
                            "title": "Resume",
                        },
                        {
                            "type": "text",
                            "text": ANALYSIS_PROMPT.format(job_description=job_description),
                        },
                    ],
                }
            ]
        elif filename.endswith(".docx"):
            resume_text = extract_text_from_docx(content)
            if not resume_text.strip():
                raise HTTPException(
                    status_code=400, detail="Could not extract text from the DOCX file."
                )
            messages = [
                {
                    "role": "user",
                    "content": (
                        f"RESUME CONTENT:\n{resume_text}\n\n"
                        + ANALYSIS_PROMPT.format(job_description=job_description)
                    ),
                }
            ]
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload a PDF or DOCX file.",
            )

        print(f"[analyze] Sending request to Claude for: {filename}")
        response = client.messages.create(
            model="claude-sonnet-4-6",   # 5× cheaper than Opus, same quality for structured JSON
            max_tokens=8000,             # enough headroom so JSON is never truncated
            messages=messages,
        )

        # Log token usage so user can track cost
        usage = response.usage
        input_tokens  = usage.input_tokens
        output_tokens = usage.output_tokens
        # Sonnet pricing: $3/M input, $15/M output
        cost_usd = (input_tokens * 3 + output_tokens * 15) / 1_000_000
        print(
            f"[analyze] Done — input={input_tokens} output={output_tokens} "
            f"est_cost=${cost_usd:.4f}"
        )

        result_text = next(
            (block.text for block in response.content if block.type == "text"), ""
        )

        # Try strict JSON extraction first, then fall back to looser parse
        json_match = re.search(r"\{.*\}", result_text, re.DOTALL)
        if not json_match:
            print(f"[analyze] Raw model output (no JSON found):\n{result_text[:800]}")
            raise HTTPException(
                status_code=500, detail="Model returned an unparseable response."
            )

        raw_json = json_match.group()
        try:
            result = json.loads(raw_json)
        except json.JSONDecodeError:
            # Truncated JSON — try to recover by trimming to last complete field
            print(f"[analyze] JSON decode failed, raw tail:\n{raw_json[-300:]}")
            raise HTTPException(
                status_code=500,
                detail="Response was cut off mid-JSON. This is rare — please try again."
            )

        # Attach token/cost info to response
        result["_usage"] = {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "est_cost_usd": round(cost_usd, 4),
        }

        # Persist to history
        row_id = save_analysis(resume.filename or "resume", job_description, result)
        result["history_id"] = row_id

        return JSONResponse(result)

    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="Failed to parse analysis JSON.")
    except anthropic.APIError as e:
        print(f"[analyze] Anthropic API error: {e}")
        raise HTTPException(status_code=500, detail=f"Anthropic API error: {e}")
    except Exception as e:
        print(f"[analyze] Unexpected error: {e}")
        raise HTTPException(status_code=500, detail=f"Server error: {e}")


# Serve the frontend (must be last)
app.mount("/", StaticFiles(directory=STATIC_DIR, html=True), name="static")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=False)
